from rest_framework.views import APIView, View
from rest_framework import viewsets
from rest_framework import status
from rest_framework_simplejwt.tokens import RefreshToken, format_lazy
from rest_framework.response import Response
from rest_framework.permissions import AllowAny
from rest_framework.generics import ListAPIView, CreateAPIView
from django.contrib.auth import authenticate
from django.contrib.auth import login as auth_login

from django.contrib.auth import get_user_model
from .serializers import (
    UserRegisterSerializer,
    UsersListSerializer,
    LoginSerializer,
    UserProtoSerializer,
    StrokeDataUploadSerializer,
)
from .models import EmailPhoneUser, Device
from .utils import *
from .mixins import *
from .parsers import *
from .services import *
from django.utils import timezone
from django.contrib import messages
from .models import Subscriber
from django.db import IntegrityError
from .tasks import async_send_newsletter
from django.shortcuts import redirect
from .mixins import DeviceMixin, LoginThrottlingMixin
import requests
from .forms import WordForm
from django.core.files.storage import default_storage

# Create your views here.


# Magic Link
from django.conf import settings
import jwt
from django.core.mail import send_mail
from django.contrib.auth import authenticate, login
from django.contrib.auth.models import User
from django.http import HttpResponse
from .forms import SignUpForm
from django.urls import reverse


class UsersListView(ListAPIView):
    permission_classes = [
        AllowAny,
    ]
    serializer_class = UsersListSerializer
    queryset = get_user_model().objects.all()

    def get(self, request):
        users = get_user_model().objects.all()
        serializer = UsersListSerializer(users, many=True)
        return Response(serializer.data, status=status.HTTP_200_OK)


M3O_API_TOKEN = "ODU3MzFmN2ItNjM3Ny00Y2MzLWEzNjktMjMxNGE5MjU1MmZl"

url = "https://api.m3o.com/v1/user/Create"
headers = {
    "Content-Type": "application/json",
    "Authorization": f"Bearer {M3O_API_TOKEN}",
}


class UserRegisterView(APIView):
    permission_classes = [
        AllowAny,
    ]
    serializer_class = UserRegisterSerializer
    queryset = get_user_model().objects.all()

    def perform_create(self, serializer):
        return serializer.save()

    def post(self, request, *args, **kwargs):
        serializer = self.serializer_class(data=request.data)
        if serializer.is_valid(raise_exception=True):
            user = self.perform_create(serializer)
            user.set_password(request.POST.get("password"))

            # add tokens
            refresh = RefreshToken.for_user(user)
            res = {
                "refresh": str(refresh),
                "access": str(refresh.access_token),
            }

            data = {
                "email": f'{serializer.validated_data["username"]}',
                "id": f"{user.id}",
                "password": f'{request.POST.get("password")}',
                "username": f'{serializer.validated_data["username"]}',
            }
            response = requests.post(url, headers=headers, json=data)
            print(response.status_code)
            print(response.json())
            print(data)

            return Response(res, status=status.HTTP_201_CREATED)


class LoginView(APIView):
    permission_classes = [
        AllowAny,
    ]
    serializer_class = LoginSerializer

    def post(self, request, *args, **kwargs):
        serializer = self.serializer_class(
            context={"request": request}, data=request.data
        )
        if serializer.is_valid(raise_exception=True):
            username = serializer.validated_data["username"]
            password = request.POST.get("password")

            user_login = authenticate(
                username=username,
                password=password,
                backend="phone.authenticate.EmailModelBackend",
            ) or authenticate(
                username=username,
                password=password,
                backend="phone.authenticate.PhoneModelBackend",
            )

            refresh = RefreshToken.for_user(user_login)
            ip_address = get_ip_address(request)
            device_name, device_details = get_device_details(
                request.META, str(refresh.access_token)
            )
            device = Device.objects.create(
                user=user_login,
                last_request_datetime=timezone.now(),
                name=device_name,
                details=device_details,
                permanent_token=refresh,
                ip_address=ip_address,
            )
            device.save()

            auth_login(request, user_login)
            user = get_user_model().objects.get(username=username)
            user_device = Device.objects.filter(user=user).first()
            # print(user.device)
            if (
                user_device.ip_address != ip_address
                or user_device.device.name != device_name
            ):
                warning_mail_send(username, ip_address)

            res = {
                "refresh": str(refresh),
                "access": str(refresh.access_token),
            }
            return Response(res, status=status.HTTP_200_OK)


# Login with Mixin


class LoginView2(APIView, DeviceMixin):
    permission_classes = [
        AllowAny,
    ]
    serializer_class = LoginSerializer

    def post(self, request, *args, **kwargs):
        serializer = self.serializer_class(
            context={"request": request}, data=request.data
        )
        if serializer.is_valid(raise_exception=True):
            username = serializer.validated_data["username"]
            password = request.POST.get("password")

            user_login = authenticate(
                username=username,
                password=password,
                backend="phone.authenticate.EmailModelBackend",
            ) or authenticate(
                username=username,
                password=password,
                backend="phone.authenticate.PhoneModelBackend",
            )

            refresh = RefreshToken.for_user(user_login)

            res = {
                "refresh": str(refresh),
                "access": str(refresh.access_token),
            }
            return Response(res, status=status.HTTP_200_OK)


def subscribe(request):
    if request.method == "POST":
        email = request.POST.get("email")
        subscriber = Subscriber(email=email, confirmed=True)
        if subscriber == Subscriber.objects.filter(email=subscriber.email):
            messages.error(request, "You are already subscribed to our newsletter!")
            return redirect("login")
        else:
            try:
                subscriber.save()
                async_send_newsletter.delay()
                messages.success(request, "You have been subscribed to our newsletter!")
                return redirect("login")
            except IntegrityError as e:
                messages.error(request, "You are already subscribed to our newsletter!")
                return redirect("login")
    else:
        return redirect("login")


def unsubscribe(request):
    confirme_subscribers = Subscriber.objects.get(email=request.GET["email"])
    for subscriber in confirme_subscribers:
        subscriber.delete()
        messages.success(
            request, "You have successfully unsubscribed from our newsletter!"
        )
        return redirect("login")


from django.conf import settings
from django.shortcuts import render


def get_definition(request):
    if request.method == "POST":
        form = WordForm(request.POST)
        if form.is_valid():
            api_key = settings.DICTIONARY_KEY
            endpoint = "https://www.dictionaryapi.com/api/v3/references/thesaurus/json/"
            word = form.cleaned_data["input"]
            url = f"{endpoint}{word}?key={api_key}"
            response = requests.get(url)
            data = response.json()
            # data = JsonResponse(response, safe=False)
            synonyms = data[0]["meta"]["syns"][0]
            meaning = data[0]["shortdef"][0]
            return render(
                request,
                "dictionary.html",
                {"synonyms": synonyms, "meaning": meaning, "word": word},
            )
    else:
        form = WordForm()
    return render(request, "dictionary.html", {"form": form})


# https://www.dictionaryapi.com/api/v3/references/thesaurus/json/test?key=15b1025e-5041-4d7b-9f7b-4f74bd0deabe


from .models import Notification


def notifications(request):
    if request.method == "POST":
        # Create a new notification object
        title = request.POST.get("title")
        body = request.POST.get("body")
        notification = Notification.objects.create(title=title, body=body)
        notification.save()

    # Get all notification objects
    notifications = Notification.objects.all()
    return render(request, "index.html", {"notifications": notifications})


# class StrokeDataUpload(APIView, DataValidationMixin):
#     permission_classes = [
#         AllowAny,
#     ]
#     parser_classes = [WordParser, PDFParser, JSONParser, CSVParser, ExcelParser]
#     serializer_class = StrokeDataUploadSerializer

#     def post(self, request):
#         # Create a serializer instance with the request data
#         serializer = StrokeDataUploadSerializer(data=request.FILES)
#         # Validate the serializer
#         serializer.is_valid(raise_exception=True)
#         # Parse the file using the `parse` method of the FileTypeMixin
#         data = self.parse(serializer.validated_data["stroke"], request.content_type)
#         # Validate the data using the DataValidationMixin
#         self.validate_data(data)
#         # Save the serializer data to the Stroke model
#         # Stroke.objects.create(**data)
#         print("data", data)

#         return Response(
#             {"msg": "Stroke data uploaded successfully"}, status=status.HTTP_201_CREATED
#         )

# from rest_framework.parsers import FileUploadParser
from rest_framework import generics


class StrokeDataUpload(generics.CreateAPIView):
    serializer_class = StrokeDataUploadSerializer
    # parser_classes = [FileUploadParser]

    parser_classes = [WordParser, PDFParser, JSONParser, CSVParser, ExcelParser]

    def post(self, request, filename, format=None, *args, **kwargs):
        print(request.data)
        serializer = self.serializer_class(data=request.FILES)
        if serializer.is_valid():
            stroke = serializer.validated_data["stroke"]
            print(str(stroke))
            print(request.content_type)
            # Do something with the uploaded file
            return Response(status=200)


from rest_framework import generics


# class FileUploadView(DataValidationMixin, generics.CreateAPIView):
#     serializer_class = StrokeDataUploadSerializer
#     content_type_whitelist = [
#         "text/csv",
#         "application/json",
#         "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#         "application/pdf",
#         "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
#     ]

#     def post(self, request, format=None, *args, **kwargs):
#         serializer = self.get_serializer(data=request.data)
#         serializer.is_valid(raise_exception=True)
#         uploaded_file = request.FILES["stroke"]
#         content_type = uploaded_file.content_type
#         print("content_type", content_type)
#         if content_type not in self.content_type_whitelist:
#             return Response(
#                 {"error": "Invalid file type"}, status=status.HTTP_400_BAD_REQUEST
#             )

#         print("uploaded_file", uploaded_file)
#         if content_type == "text/csv":
#             import csv

#             with open(uploaded_file.name, "r", encoding="utf-8") as f:
#                 reader = csv.DictReader(f)
#                 contents = [x for x in reader]

#                 self.validate_data(request, data=contents[0].keys())

#         if (
#             content_type
#             == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
#         ):
#             import openpyxl

#             workbook = openpyxl.load_workbook(uploaded_file.name)
#             sheet_names = workbook.sheetnames
#             sheet = workbook[sheet_names[0]]
#             # Print the sheet data
#             for row in sheet.rows:
#                 for cell in row:
#                     print(cell.value)

#         if (
#             content_type
#             == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#         ):
#             import docx

#             document = docx.Document(uploaded_file.name)
#             print(document.paragraphs[0].text)

#         # Now you can use the file as needed, such as by saving it to a database or storing it on a file system.

#         return Response(
#             {"message": "Saved successfully"}, status=status.HTTP_201_CREATED
#         )

from .models import Stroke


class FileUploadView(DataValidationMixin, generics.CreateAPIView):
    serializer_class = StrokeDataUploadSerializer
    content_type_whitelist = [
        "text/csv",
        "application/json",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ]

    def post(self, request, format=None, *args, **kwargs):
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        validated_data = serializer.validated_data
        uploaded_file = request.FILES["stroke"]
        file_name = uploaded_file.name
        content_type = uploaded_file.content_type
        if content_type not in self.content_type_whitelist:
            return Response(
                {"error": "Invalid file type"}, status=status.HTTP_400_BAD_REQUEST
            )
        try:
            if content_type == "text/csv":
                # import csv

                # with open(uploaded_file.name, "r", encoding="utf-8") as f:
                #     reader = csv.DictReader(f)
                #     contents = [x for x in reader]
                contents = read_csv_file(uploaded_file)
                self.validate_data(request, data=contents[0].keys())
                # print("contents", contents[0])
                # Create a new Stroke instance with the data in the row
                # default_storage.save(
                #     "users_uploaded_files/{0}/{1}".format(
                #         request.user.username, file_name
                #     )
                # )

                stroke = Stroke(
                    user=request.user,
                    glucose_level=contents[0]["glucose_level"],
                    bmi=contents[0]["bmi"],
                    blood_pressure_systolic=contents[0]["blood_pressure_systolic"],
                    blood_pressure_diastolic=contents[0]["blood_pressure_diastolic"],
                    age=contents[0]["age"],
                )
                stroke.save()
                # stroke = Stroke(
                #     user=request.user,
                #     glucose_level=row["glucose_level"],
                #     bmi=row["bmi"],
                #     blood_pressure_systolic=row["blood_pressure_systolic"],
                #     blood_pressure_diastolic=row["blood_pressure_diastolic"],
                #     age=row["age"],
                # )
                # stroke.save()
                print("saving stroke data")
            if (
                content_type
                == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            ):
                import openpyxl

                workbook = openpyxl.load_workbook(uploaded_file.name)
                sheet_names = workbook.sheetnames
                sheet = workbook[sheet_names[0]]
            if (
                content_type
                == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ):
                import docx

                document = docx.Document(uploaded_file.name)
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_400_BAD_REQUEST)
        return Response(
            {"message": "Saved successfully"}, status=status.HTTP_201_CREATED
        )


# views.py
from django.http import HttpResponse
import datetime


def time(request):
    now = datetime.datetime.now()
    return HttpResponse(now.strftime("%Y-%m-%d %H:%M:%S"))


from django.shortcuts import render

time = 0


def time_view(request):
    global time
    context = {"time": time}
    return render(request, "time.html", context)


def start_view(request):
    global time
    if request.method == "POST":
        # process form data
        time += 1
        context = {"time": time}
        return render(request, "time.html", context)
    else:
        # render the form template
        return render(request, "start.html")


def reset_view(request):
    global time
    if request.method == "POST":
        # process form data
        time = 0
        context = {"time": time}
        return render(request, "time.html", context)
    else:
        # render the form template
        return render(request, "reset.html")


from django.shortcuts import render
import datetime
from django.views.decorators.csrf import csrf_exempt


def counter(request):
    current_time = datetime.datetime.now()
    return render(request, "counter.html", {"current_time": current_time})


@csrf_exempt
def stop(request):
    return render(request, "counter.html", {"current_time": None})


# @csrf_exempt
# def stopwatch(request):
#     elapsed_time = 0  # replace this with the elapsed time from the database
#     return render(request, 'stopwatch.html', {'elapsed_time': elapsed_time})


import random


@csrf_exempt
def stopwatch(request):
    elapsed_time = 0  # replace this with the elapsed time from the database
    # Select a random picture from a list of picture URLs
    picture_urls = [
        "https://picsum.photos/200/300",
        "https://picsum.photos/200/301",
        "https://picsum.photos/200/302",
        "https://picsum.photos/200/303",
    ]
    random_picture_url = random.choice(picture_urls)
    return render(
        request,
        "stopwatch.html",
        {"elapsed_time": elapsed_time, "random_picture_url": random_picture_url},
    )


from .serializers import SubscriptionSerializer
from .models import Subscription


class CreateSubscriptionView(generics.CreateAPIView):
    queryset = Subscription.objects.all()
    serializer_class = SubscriptionSerializer


from django.http import JsonResponse


class PaymentWebhookView(View):
    def post(self, request):
        # Extract the payment information from the request
        payment_info = request.POST.get("payment_info")

        # Use the payment information to create or update a subscription
        subscription = process_payment(payment_info)

        return JsonResponse({"status": "success"})


def process_payment(payment_info):
    # Extract necessary information from the payment_info
    subscription_id = payment_info["subscription_id"]
    amount = payment_info["amount"]
    start_date = payment_info["start_date"]
    end_date = payment_info["end_date"]
    frequency = payment_info["frequency"]
    next_delivery_date = payment_info["next_delivery_date"]

    try:
        # check if a subscription with the given id already exists
        subscription = Subscription.objects.get(id=subscription_id)
        # Update the subscription based on the payment information
        subscription.amount = amount
        subscription.status = status
        subscription.save()
        return subscription
    except Subscription.DoesNotExist:
        # Create a new subscription
        subscription = Subscription.objects.create(
            id=subscription_id,
            amount=amount,
            start_date=start_date,
            end_date=end_date,
            frequency=frequency,
            next_delivery_date=next_delivery_date,
        )
        return subscription


from django.shortcuts import render
import requests


def search_movie(request):
    if request.method == "POST":
        movie_title = request.POST["movie_title"]
        api = "6e907d6d"
        url = f"http://www.omdbapi.com/?apikey={api}&t={movie_title}"
        response = requests.get(url)
        data = response.json()
        title = data["Title"]
        year = data["Year"]
        rating = data["Ratings"][0]["Value"]
        release_date = data["Released"]
        language = data["Language"]
        poster = data["Poster"]

        movie = {
            "Title": title,
            "Year": year,
            "Rating": rating,
            "Released": release_date,
            "Language": language,
            "Poster": poster,
        }
        print(movie)
        return render(request, "start.html", {"movies": movie})
    else:
        return render(request, "start.html")


# views.py
from django.shortcuts import render
from .forms import Page1Form, Page2Form, Page3Form


def forms(request):
    if request.method == "POST":
        if "page1" in request.POST:
            form = Page2Form()
            return render(request, "page2.html", {"form": form})
        elif "page2" in request.POST:
            form = Page3Form()
            return render(request, "page3.html", {"form": form})
        else:
            # save the form data
            # ...
            return render(request, "completed.html")
    else:
        form = Page1Form()
    return render(request, "forms.html", {"form": form})


from .forms import MoveForm


def make_move(request, game_id):
    game = get_object_or_404(Game, pk=game_id)
    board = [list(row) for row in game.board.split("-")]
    player = game.next_player
    form = MoveForm(request.POST or None, board=board)
    if form.is_valid():
        row = form.cleaned_data["row"]
        col = form.cleaned_data["col"]
        board[row][col] = player.symbol
        game.board = "-".join("".join(row) for row in board)
        game.next_player = game.other_player
        if game.is_finished():
            game.winner = player
            game.finished = True
        game.save()
        return redirect("game_detail", game_id=game.id)
    return render(request, "game_detail.html", {"game": game, "form": form})


from django.shortcuts import render, get_object_or_404
from .models import Game, Player


def game_detail(request, game_id):
    game = get_object_or_404(Game, pk=game_id)
    return render(request, "game_detail.html", {"game": game})


def game_list(request):
    games = Game.objects.all()
    return render(request, "game_list.html", {"games": games})


def player_detail(request, player_id):
    player = get_object_or_404(Player, pk=player_id)
    return render(request, "player_detail.html", {"player": player})


def player_list(request):
    players = Player.objects.all()
    return render(request, "player_list.html", {"players": players})


from rest_framework import status
from rest_framework.generics import GenericAPIView
from rest_framework.response import Response
from django.contrib.auth.tokens import default_token_generator
from django.contrib.sites.shortcuts import get_current_site
from django.utils.encoding import force_bytes, force_str
from django.utils.http import urlsafe_base64_encode, urlsafe_base64_decode
from django.core.mail import send_mail

from .serializers import PasswordResetSerializer, PasswordResetConfirmSerializer


class PasswordResetView(GenericAPIView):
    serializer_class = PasswordResetSerializer

    def post(self, request, *args, **kwargs):
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        email = serializer.validated_data["email"]
        user = get_user_model().objects.get(email=email)
        current_site = get_current_site(request)
        uid = urlsafe_base64_encode(force_bytes(user.pk))
        token = default_token_generator.make_token(user)
        password_reset_url = f"{current_site}/password-reset-confirm/{uid}/{token}"
        email_subject = "Password Reset Request"
        email_body = (
            f"Please follow the link to reset your password: {password_reset_url}"
        )
        send_mail(
            email_subject, email_body, "from@example.com", [email], fail_silently=False
        )
        return Response({"email": email}, status=status.HTTP_200_OK)


class PasswordResetConfirmView(GenericAPIView):
    serializer_class = PasswordResetConfirmSerializer

    def post(self, request, uidb64, token, *args, **kwargs):
        try:
            uid = force_str(urlsafe_base64_decode(uidb64))
            user = get_user_model().objects.get(pk=uid)
        except (TypeError, ValueError, OverflowError, User.DoesNotExist):
            user = None

        if user is not None and default_token_generator.check_token(user, token):
            serializer = self.get_serializer(data=request.data)
            serializer.is_valid(raise_exception=True)
            user.set_password(serializer.validated_data["new_password"])
            user.save()
            return Response(
                {"status": "password reset successful"}, status=status.HTTP_200_OK
            )
        return Response(
            {"status": "password reset failed"}, status=status.HTTP_400_BAD_REQUEST
        )


from django.http import JsonResponse
from .models import DataPoint


def get_data(request):
    data = list(DataPoint.objects.values("value", "created_at"))
    return JsonResponse({"data": data})
    # return render(request, 'chart.html', {'data': data})


from django.shortcuts import render


def chart_view(request):
    return render(request, "chart.html")


@csrf_exempt
def add_data(request):
    if request.method == "POST":
        value = request.POST.get("value")
        DataPoint.objects.create(value=value)
        return JsonResponse({"status": "success"})


from django.shortcuts import render, redirect
from .forms import SignUpForm


def signup(request):
    if request.method == "POST":
        form = SignUpForm(request.POST)
        if form.is_valid():
            # process the form data
            return redirect("success")

    else:
        form = SignUpForm()
    return render(request, "signup2.html", {"form": form})


def success(request):
    return render(request, "success.html")


import openai

openai.api_key = "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"

from .forms import CodeForm


def code_upload(request):
    if request.method == "POST":
        form = CodeForm(request.POST, request.FILES)
        if form.is_valid():
            code_text = form.cleaned_data.get("code_text")
            code_file = request.FILES.get("code_file")
            if code_text:
                # process the code text
                prompt = "Fix the following code:\n" + code_text
                response = (
                    openai.Completion.create(
                        engine="text-davinci-002",
                        prompt=prompt,
                        max_tokens=1024,
                        n=1,
                        stop=None,
                        temperature=0.5,
                    )
                    .choices[0]
                    .text
                )
                # ...
                return render(request, "success.html", {"response": response})
            elif code_file:
                # process the code file
                code_text = code_file.read().decode("utf-8")
                prompt = "Fix the following code:\n" + code_text
                response = (
                    openai.Completion.create(
                        engine="text-davinci-002",
                        prompt=prompt,
                        max_tokens=1024,
                        n=1,
                        stop=None,
                        temperature=0.5,
                    )
                    .choices[0]
                    .text
                )
                # ...
                return render(request, "success.html", {"response": response})
            else:
                # handle the case where neither code text nor code file is provided
                # ...
                return render(request, "error.html")
    else:
        form = CodeForm()
    return render(request, "code_upload.html", {"form": form})


# views.py
import asyncio
from channels.db import database_sync_to_async
from channels.layers import get_channel_layer
from django.shortcuts import get_object_or_404
from django.http import JsonResponse
from django.views.decorators.http import require_http_methods
from .models import Post
from channels.layers import get_channel_layer
from asgiref.sync import async_to_sync


@csrf_exempt
@require_http_methods(["POST"])
def like_post_view(request, post_id):
    post = get_object_or_404(Post, id=post_id)
    post.like_count += 1
    post.save()

    async def send_like_update():
        channel_layer = get_channel_layer()
        group_name = "test_group"
        print("Sending update to group {}".format(group_name))
        await channel_layer.group_send(
            group_name,
            {"type": "update_post_likes", "post_id": post_id},
        )
        print("Update sent")

    async_to_sync(send_like_update)()

    return JsonResponse({"success": post.like_count})


def post_list(request):
    posts = Post.objects.all()
    return render(request, "template.html", {"posts": posts})


def list(request):
    return render(request, "list.html")


from .forms import FormStep1, FormStep2


def form_step1(request):
    if request.method == "POST":
        form = FormStep1(request.POST)
        form.name = request.POST.get("name")
        form.email = request.POST.get("email")
        if form.is_valid():
            form.save()
            return redirect("form_step2")
    else:
        form = FormStep1()
    return render(request, "form_step1.html", {"form": form})


def form_step2(request):
    if request.method == "POST":
        form = FormStep2(request.POST)
        form.phone = request.POST.get("phone")
        form.address = request.POST.get("address")
        if form.is_valid():
            form.save()
            return redirect("success")
    else:
        form = FormStep2()
    return render(request, "form_step2.html", {"form": form})


from django.shortcuts import render, redirect
from .forms import MyForm
from .models import MyModel

from django.shortcuts import render, redirect
from .forms import MyForm


def my_view(request):
    choices = MyModel.objects.all().values_list("name", flat=True)
    if request.method == "POST":
        form = MyForm(request.POST)
        if form.is_valid():
            # Process the form data as before
            # ...
            return redirect("success")
    else:
        form = MyForm()
    return render(request, "my_template.html", {"form": form, "choices": choices})


from django.http import JsonResponse
from django.views.decorators.http import require_GET
from .models import MyModel


@require_GET
def get_choices(request):
    # Query the choices from the database
    choices = list(MyModel.objects.values_list("name", flat=True))
    return JsonResponse(choices, safe=False)


from django.shortcuts import render
from django.http import HttpResponseRedirect
from django.urls import reverse
from django.contrib import messages

import phonenumbers
from phonenumbers import geocoder, carrier

from .forms import ContactForm


def contact(request):
    if request.method == "POST":
        form = ContactForm(request.POST)
        if form.is_valid():
            # Get the cleaned data from the form
            name = form.cleaned_data["name"]
            email = form.cleaned_data["email"]
            country = form.cleaned_data["country"]
            phone_number = form.cleaned_data["phone"]
            message = form.cleaned_data["message"]

            # Parse the phone number using phonenumbers
            parsed_number = phonenumbers.parse(phone_number, country)

            # Add the country code to the phone number
            phone_with_country_code = phonenumbers.format_number(
                parsed_number, phonenumbers.PhoneNumberFormat.E164
            )

            # Get the carrier and geographic information for the phone number
            carrier_name = carrier.name_for_number(parsed_number, "en")
            location = geocoder.description_for_number(parsed_number, "en")

            # TODO: Do something with the contact form data, e.g. send an email or save to a database

            # Show a success message and redirect to the home page
            messages.success(request, "Thank you for your message!")
            return HttpResponseRedirect(reverse("home"))
    else:
        form = ContactForm()

    context = {
        "form": form,
    }
    return render(request, "contact.html", context)



from django.contrib.auth.views import LoginView
from django.contrib.auth import authenticate, login
from django.http import HttpResponseRedirect
from datetime import timedelta
from django.urls import reverse_lazy
from .forms import RememberMeAuthenticationForm


class CustomLoginView(LoginView):
    template_name = 'remember.html'
    form_class = RememberMeAuthenticationForm
    success_url = reverse_lazy('chart_view') 


      
    def get_success_url(self):
        return self.success_url

    def form_valid(self, form):
        remember_me = form.cleaned_data.get('remember_me')
        if remember_me:
            print("yes it working")
            self.request.session.set_expiry(timedelta(days=30))
        else:
            self.request.session.set_expiry(0)
        return super().form_valid(form)



from django.contrib.auth.views import LogoutView
from django.http import HttpResponseRedirect

class CustomLogoutView(LogoutView):
    next_page = '/remember/'

    def dispatch(self, request, *args, **kwargs):
        response = super().dispatch(request, *args, **kwargs)
        if request.user.is_authenticated:
            request.session.flush()
            response.delete_cookie('sessionid')
        return response


from supabase import create_client
import supabase
from django.shortcuts import redirect
from django.urls import reverse

SUPABASE_URL = 'https://ccfmytlvtfuhqhucbeuu.supabase.co'
SUPABASE_KEY = 'eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6ImNjZm15dGx2dGZ1aHFodWNiZXV1Iiwicm9sZSI6ImFub24iLCJpYXQiOjE2NzczNzQxMTgsImV4cCI6MTk5Mjk1MDExOH0.e63wPAN2hbaoKgSLtFvotZsS5bhG9uM-NBhawbbaPj0'

from supabase import create_client
 
# supabase login

def login_supabase(request):
    if request.method == "POST":
        form = SignUpForm(request.POST)
        if form.is_valid():
            username = form.cleaned_data["username"]
            password = form.cleaned_data["password"]
            client = create_client(SUPABASE_URL, SUPABASE_KEY)
            try:
                user = client.query(
                    "SELECT * FROM users WHERE username = :username AND password = :password",
                    {"username": username, "password": password},
                )
                if user:
                    user = user[0]
                    request.session["user_id"] = user["id"]
                    request.session["username"] = user["username"]
                    return redirect("home")
                else:
                    form.add_error("username", "Invalid username or password")
            except SupabaseError as e:
                form.add_error("username", str(e))
    else:
        form = SignUpForm()
    return render(request, "login.html", {"form": form})
     
SUPABASE_URL = 'https://ccfmytlvtfuhqhucbeuu.supabase.co'
SUPABASE_KEY = 'eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6ImNjZm15dGx2dGZ1aHFodWNiZXV1Iiwicm9sZSI6ImFub24iLCJpYXQiOjE2NzczNzQxMTgsImV4cCI6MTk5Mjk1MDExOH0.e63wPAN2hbaoKgSLtFvotZsS5bhG9uM-NBhawbbaPj0'

client = create_client(SUPABASE_URL, SUPABASE_KEY)
auth = client.auth



# def login(request):
#     supabase = create_client(SUPABASE_URL, SUPABASE_KEY)
#     redirect_url = request.build_absolute_uri(reverse('callback'))
#     url = auth.authorization_url(redirect_url)
#     return redirect(url)
    

 
# def callback(request):
#     supabase = create_client(SUPABASE_URL, SUPABASE_KEY)
#     redirect_url = request.build_absolute_uri(reverse('callback'))
#     code = request.GET.get('code')
#     if not code:
#         return redirect(reverse('login'))
#     token = supabase.auth.exchange_code_for_token(code, redirect_url)
#     request.session['access_token'] = token['access_token']
#     return redirect(reverse('home'))
     



# from django.shortcuts import redirect
# from django.urls import reverse
# from django.views.generic import View
# from django.conf import settings

# class LoginView(View):
#     print("inside login view")
#     def get(self, request):
#         # Generate the Supabase authentication URL
#         redirect_uri = request.build_absolute_uri(reverse('callback'))
#         print(redirect_uri)
#         auth_url = f"{settings.SUPABASE_URL}/oauth/v2/auth?client_id={settings.SUPABASE_KEY}&redirect_uri={redirect_uri}&response_type=code&scope=openid%20profile%20email"

#         # Redirect the user to the authentication URL
#         return redirect(auth_url)



from django.shortcuts import redirect, reverse
from django.views import View
from supabase import create_client



from django.shortcuts import redirect, reverse
from django.views import View
from supabase import create_client


class LoginView(View):
    def get(self, request):
        client = create_client(SUPABASE_URL, SUPABASE_KEY)
        auth = client.auth

        # generate the authorization URL to start the authentication flow
        auth_url = auth.authorization_url()

        # store the auth state in the user's session
        request.session['supabase_auth_state'] = auth.auth_state_token

        # redirect the user to the authorization URL
        return redirect(auth_url)

    def post(self, request):
        client = create_client(SUPABASE_URL, SUPABASE_KEY)
        auth = client.auth

        # extract the access token and refresh token from the user's session
        access_token = request.session.get('supabase_access_token')
        refresh_token = request.session.get('supabase_refresh_token')

        if not access_token or not refresh_token:
            # access token or refresh token is missing, redirect back to the login page
            return redirect(reverse('login'))

        # authenticate the user with Supabase using the access token and refresh token
        auth.set_access_token(access_token)
        auth.set_refresh_token(refresh_token)
        user = auth.user()

        # check if the user is authenticated
        if not user:
            # user is not authenticated, redirect back to the login page
            return redirect(reverse('login'))

        # user is authenticated, log them in to your Django app
        # ...

        # redirect the user to the home page
        return redirect(reverse('home'))


def supabase_callback(request):
    client = create_client(SUPABASE_URL, SUPABASE_KEY)
    auth = client.auth

    # extract the authorization code and state token from the request
    code = request.GET.get('code')
    state = request.GET.get('state')

    # check that the state token matches the one stored in the user's session
    if state != request.session.get('supabase_auth_state'):
        # state token doesn't match, redirect back to the login page
        return redirect(reverse('login'))

    # exchange the authorization code for an access token and refresh token
    token_response = auth.exchange_code_for_token(code)

    # store the access token and refresh token in the user's session
    request.session['supabase_access_token'] = token_response['access_token']
    request.session['supabase_refresh_token'] = token_response['refresh_token']

    # redirect the user to the home page
    return redirect(reverse('success'))



from .models import Book

def book_list(request):
    books = Book.objects.all()
    return render(request, 'book_list.html', {'books': books})



from django.shortcuts import render
from django.http import JsonResponse
import zxcvbn

# @csrf_exempt
# def validate_password_strength(request):
#     password = request.POST.get('password')
#     result = zxcvbn.zxcvbn(password)
#     score = result['score']
#     if score < 3:
#         strength = 'Weak'
#     elif score < 4:
#         strength = 'Fair'
#     else:
#         strength = 'Strong'
#     return JsonResponse({'strength': strength})




from django.shortcuts import render
import json
import re
from zxcvbn import zxcvbn

def password_strength(request):
        return render(request, 'password.html')




# @csrf_exempt
# def validate_password_strength(request):
#     if request.method == 'POST':
#         password = request.POST.get('password')
#         score = calculate_password_strength(password)
#         print(password)
#         if score >= 10:
#             strength = 'Strong'
#             return JsonResponse({'strength': strength})

#         elif score >= 5:
#             strength = 'Fair'
#             return JsonResponse({'strength': strength})

#         else:
#             strength = 'Weak'
#             return JsonResponse({'strength': strength})

# def calculate_password_strength(password):
#     score = 0
#     if len(password) < 8:
#         return score
#     score += 1
#     if re.search(r'[a-z]', password):
#         score += 1
#     if re.search(r'[A-Z]', password):
#         score += 1
#     if re.search(r'[0-9]', password):
#         score += 1
#     if re.search(r'[!@#$%^&*(),.?":{}|<>]', password):
#         score += 1
#     return score

# @csrf_exempt
# def validate_password_strength(request):
#     if request.method == 'POST':
#         password = request.POST.get('password', '')
#         if password:
#             # Calculate password strength score using zxcvbn library
#             score = zxcvbn.password_strength(password)['score']
#             if score > 8:
#                 strength = 'Excellent'
#             elif score > 5:
#                 strength = 'Strong'
#             elif score > 3:
#                 strength = 'Fair'
#             else:
#                 strength = 'Weak'
#             return JsonResponse({'strength': strength})
#     return JsonResponse({'error': 'Invalid request method'})



@csrf_exempt
def validate_password_strength(request):
    if request.method == 'POST':
        password = request.POST.get('password', '')
        result = zxcvbn(password)
        score = result['score']
        if score == 2:
            return JsonResponse({'strength': 'Strong', 'score': score})
        elif score == 1:
            return JsonResponse({'strength': 'Fair', 'score': score})
        elif score == 3:
            return JsonResponse({'strength': 'Very Strong', 'score': score})

        elif score == 4:
            return JsonResponse({'strength': 'Excellent', 'score': score})
        elif score == 0:
            return JsonResponse({'strength': 'Weak', 'score': score})




from django.http import HttpResponse
from . import rabbitmq

def index(request):
    rabbitmq.channel.basic_publish(
        exchange='',
        routing_key='hello',
        body='Hello World!'
    )
    return HttpResponse("Message sent to RabbitMQ")



from django.contrib.auth import authenticate, login
from django.core.mail import send_mail
from django.conf import settings
from .forms import SignUpForm1

def signup1(request):
    if request.method == 'POST':
        form = SignUpForm1(request.POST)
        if form.is_valid():
            user = form.save(commit=False)
            user.is_active = False
            user.save()
            # Send email verification
            subject = 'Verify your email'
            message = 'Please click the link below to verify your email address: {0}'.format(settings.BASE_URL + '/verify-email/' + user.username)
            from_email = settings.DEFAULT_FROM_EMAIL
            recipient_list = [user.email]
            send_mail(subject, message, from_email, recipient_list)
            return redirect('verify-email-sent')
    else:
        form = SignUpForm1()
    return render(request, 'test/signup1.html', {'form': form})

def verify_email_sent(request):
    return render(request, 'test/verify_email_sent.html')



from django.shortcuts import render, redirect
from django.contrib.auth import authenticate, login
from django.contrib import messages
from .forms import SignUpForm2


def signup2(request, username):
    user = get_user_model().objects.get(username=username)
    if request.method == 'POST':
        form = SignUpForm2(request.POST)
        if form.is_valid():
            user.phone = form.cleaned_data['phone']
            user.set_password(form.cleaned_data['password1'])
            user.is_active = True
            user.save()
            return redirect('login')
    else:
        form = SignUpForm2()
    return render(request, 'test/signup2.html', {'form': form})



from django.contrib.auth import get_user_model
from django.core.mail import send_mail
from django.template.loader import render_to_string
from django.shortcuts import render, redirect
from django.urls import reverse
from django.conf import settings

User = get_user_model()

def resend_verification(request):
    email = request.GET.get('email')
    if email:
        try:
            user = User.objects.get(email=email)
        except User.DoesNotExist:
            user = None
        if user is not None and not user.is_verified:
            subject = 'Verify your email'
            message = render_to_string('verify_email.txt', {'user': user})
            send_mail(
                subject,
                message,
                settings.DEFAULT_FROM_EMAIL,
                [user.email],
                fail_silently=False,
            )
            return redirect(reverse('verify_email_sent') + f'?email={email}')
    return redirect('success')



def verify_email(request, username):
    user = User.objects.get(username=username)
    user.is_active = True
    user.save()
    return redirect('signup2', username=username)





from rest_framework.views import APIView
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework.response import Response
from .serializers import ImageSerializer

class ImageUploadView(APIView):
    parser_classes = (MultiPartParser, FormParser)
    serializer_class = ImageSerializer

    def post(self, request, *args, **kwargs):
        serializer = ImageSerializer(data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=201)
        else:
            return Response(serializer.errors, status=400)
